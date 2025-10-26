import streamlit as st
import pdfplumber
import re
import google.generativeai as genai
from docx import Document
from io import BytesIO
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

st.session_state.setdefault("download_ready", False)

# -----------------------------
# CONFIGURATION
# -----------------------------
GEMINI_API_KEY = "AIzaSyAKbtJyypvjhUii916BcqpAHeprwZWW3Dc"  # <-- Replace with your valid key
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# -----------------------------
# STREAMLIT FRONT END SETUP
# -----------------------------
st.set_page_config(page_title="BCBS Justification for IDR", layout="centered")
st.title("🏥 BCBS Justification for IDR Generator")
st.write("Upload EOB PDF, MRN PDF, and Prompt TXT file to generate a formatted BCBS justification document.")

# -----------------------------
# FILE UPLOADS
# -----------------------------
eob_file = st.file_uploader("📄 Upload EOB PDF", type=["pdf"])
mrn_file = st.file_uploader("🧾 Upload MRN PDF", type=["pdf"])
prompt_file = st.file_uploader("📝 Upload MRN Summary Prompt (TXT)", type=["txt"])

# -----------------------------
# HELPER FUNCTIONS
# -----------------------------
def extract_text_from_pdf(uploaded_file):
    """Extract text from both scanned and digital PDFs"""
    text = ""
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text
    except Exception:
        pass
    uploaded_file.seek(0)
    images = convert_from_bytes(uploaded_file.read())
    for img in images:
        text += pytesseract.image_to_string(img) + "\n"
    return text


def find_field(patterns, text, label):
    """Try multiple regex patterns until a match is found"""
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return f"{label} not found"


def extract_fields(eob_text):
    """Extract claim fields robustly from EOB text"""
    # ---------------- DATE OF SERVICE ----------------
    date_patterns = [
        r"Date.?of.?Service[s]?:?\s*(\d{1,2}/\d{1,2}/\d{2,4})",
        r"Service Date[s]?:?\s*(\d{1,2}/\d{1,2}/\d{2,4})",
        r"Service Dates?\s*[-–]\s*(\d{1,2}/\d{1,2}/\d{2,4})",
        r"Date Range[:\s]+(\d{1,2}/\d{1,2}/\d{2,4})"
    ]
    date_of_service = find_field(date_patterns, eob_text, "Date")

    # ---------------- HCPCS / CPT CODES ----------------
    lines = eob_text.splitlines()
    ranked_codes = []
    for line in lines:
        code_match = re.search(r"(?:HCPCS|CPT|^|\s)([789]\d{3,4})(?:\s|$)", line)
        if code_match:
            code = code_match.group(1)
            if re.search(r"[A-Za-z]{2,}\d{4,}", line):
                continue

            # Detect emergency codes (99281–99285, 99291–99292)
            if re.match(r"99(28[1-5]|29[1-2])", code):
                # Attach modifier -25 only to emergency codes
                if re.search(r"\b25\b", line):
                    code = f"{code}-25"

            if code not in ranked_codes:
                ranked_codes.append(code)

    # ---------------- DRG CODE ----------------
    drg_patterns = [
        r"DRG\s*[:#-]?\s*([0-9]{2,4})",
        r"DRG\s*Code\s*[:#-]?\s*([0-9]{2,4})",
        r"Diagnosis\s*Related\s*Group[^0-9]*([0-9]{2,4})",
        r"MS[-\s]*DRG[^0-9]*([0-9]{2,4})",
        r"DRG[^0-9]*([0-9]{2,4})",
        r"RelatedGroup[^0-9]*([0-9]{2,4})"
    ]
    drg_code = find_field(drg_patterns, eob_text, "DRG Code")

    # ---------------- BILLING PROVIDER ----------------
    billing_patterns = [
        r"Billing Provider(?: Name)?:\s*([A-Za-z0-9\s.,&'\-]+)",
        r"Billing Provider\s+([A-Za-z0-9\s.,&'\-]+)",
        r"Provider Name\s*[:\-]?\s*([A-Za-z0-9\s.,&'\-]+)"
    ]
    billing_provider = find_field(billing_patterns, eob_text, "Billing Provider")

    # Remove trailing noise
    billing_provider = re.sub(
        r"\s*(NPI.*|Other Carrier.*|Rendering Provider.*|Check Date.*|Address.*|City.*|State.*|Zip.*)$",
        "",
        billing_provider,
        flags=re.IGNORECASE,
    ).strip()

    return date_of_service, ranked_codes, drg_code, billing_provider


def generate_mrn_summary(prompt_text, mrn_text):
    """Generate structured MRN summary using Gemini"""
    genai.configure(api_key=GEMINI_API_KEY.strip())
    model = genai.GenerativeModel("gemini-2.5-flash")
    combined_prompt = (
        f"{prompt_text}\n\n---\n\n{mrn_text}\n\n"
        "Format output using **bold** headings and line breaks for clarity."
    )
    response = model.generate_content(combined_prompt)
    return response.text.strip()


def generate_bcbs_justification_letter(date, hcpcs, drg, billing_provider, mrn_summary):
    """Build the complete letter text with variables inserted."""
    hcpcs_list = ','.join(hcpcs) if hcpcs else 'N/A'
    emergency_codes = []
    for c in hcpcs:
        base_code = re.sub(r"-.*", "", c)
        if re.match(r"99(28[1-5]|29[1-2])", base_code):
            emergency_codes.append(base_code)
    emergency_code_text = emergency_codes[0] if emergency_codes else "99284"

    return f"""
This letter is submitted in support of our Independent Dispute Resolution (IDR) request under the No Surprises Act (NSA). We are challenging the reimbursement amount determined by BCBS for the emergency services rendered on **{date}**. The payment issued by BCBS does not adequately reflect the level of care provided, nor does it comply with NSA transparency requirements.

We firmly assert that a higher reimbursement is justified based on the significant medical complexity of the case and in accordance with the payment determination criteria outlined in **45 CFR §149.510(c)(4)(iii)**. This includes, but is not limited to, the acuity of the patient’s condition, the scope of services rendered, and the qualifications and experience of the attending provider.

### QPA Transparency Failure & Arbitrary Methodology
BCBS’s claims their QPA was calculated using internal, fee-for-service median contracted rates from 2020–2021 and that it excludes bonuses, shared risk, and information derived from databases. However, this explanation fails to meet the disclosure standards under 45 CFR §149.140(a)(12) and CMS NSA FAQ #12, which require the QPA to be provided on a per-service basis and not as a flat amount across unrelated CPT’s.
Furthermore, BCBS does not disclose:
•	The actual median rate or how it was derived
•	Which specialties or contracts were used
•	Whether their data reflects ghost rates, $1 floor rates, or stale agreements
These omissions contradict the purpose of the NSA’s transparency goals and materially impair providers' ability to evaluate fairness and negotiate in good faith. In effect, BCBS is asking IDR entities to accept a "black box" QPA with no line-item disclosure, no clinical justification, and no meaningful accountability.
This approach also disproportionately harms out-of-network emergency providers, especially those serving underserved populations and operating independently of hospital systems. It reinforces the need for a fair, case-specific evaluation of the actual services rendered, which far exceeds the clinical complexity of what BCBS’s rate represents.

###  Improper DRG Classification and Non-Compliant Adjudication of Outpatient Emergency Claim
The claim was submitted using CPT/HCPCS codes—**{hcpcs_list}**—each reflecting distinct, medically necessary procedures performed during the emergency department encounter. Modifier 25 was correctly appended to CPT **{emergency_code_text}**, denoting a significant, separately identifiable evaluation and management (E/M) service delivered in addition to diagnostic testing, as recognized by CMS's National Correct Coding Initiative (NCCI) policy.
Nevertheless, the BCBS EOB lists a “Hospital Payment Indicator: R – Case Rate” and assigns Diagnosis Related Group (DRG) **{drg}** with a zero DRG weight (0.00000), indicating the payer’s system improperly converted an outpatient freestanding emergency department (FSED) claim into an inpatient, DRG-based payment methodology. This reclassification is both factually incorrect and in violation of federal billing and adjudication requirements. DRG payment models are explicitly reserved for inpatient hospital stays and are not permitted as a basis for adjudicating outpatient emergency claims billed under CPT/HCPCS coding protocols.
Applying DRG automation to an FSED claim misrepresents the nature of the service, the facility type, and the context of care delivery, resulting in an unsupported case-rate payment that fails to consider the submitted codes and the true scope of services rendered. This process does not satisfy the legal requirements of 45 C.F.R. § 149.510(c)(4)(iii), which obligate payers to evaluate payment based on the provider’s experience, facility type, service scope, and patient acuity.
For these reasons, it is imperative that the adjudication and Independent Dispute Resolution (IDR) review privilege only the original CPT/HCPCS codes submitted, accurately representing the outpatient emergency care delivered. Reimbursement must be recalculated according to NSA regulations to guarantee precise, transparent, and equitable payment in line with the intent of the No Surprises Act.

### Patient Acuity & Complexity of Care
{mrn_summary}

### Training, Experience & Quality Measures
**{billing_provider}** operates as a 24/7 Freestanding Emergency Department (FSED), staffed exclusively by board-certified emergency medicine physicians and highly trained nursing professionals dedicated to delivering care that meets or exceeds nationally recognized benchmarks for clinical accuracy, quality, and patient safety. Our commitment to excellence is demonstrated through accreditation by The Joint Commission (Gold Seal of Approval), COLA’s Seal of Quality in Healthcare, and the Center for Improvement in Healthcare Quality (CIHQ)—each reflecting rigorous national compliance standards for safety, quality, and patient outcomes.
We have made significant capital and operational investments in advanced diagnostic and treatment technologies, including multi-slice CT scanners, digital radiography, and on-site laboratory services, enabling our team to deliver hospital-level emergency care efficiently while reducing the overcrowding burden typical of traditional hospital ERs. FSEDs in Texas are highly regulated under state licensure laws, undergo continuous inspections, and are bound by EMTALA-comparable obligations to ensure all patients receive appropriate emergency care regardless of insurance status or ability to pay. Independent research consistently demonstrates that Texas FSEDs deliver timely, efficient, and medically necessary care comparable to—if not exceeding—hospital-based emergency departments, particularly regarding patient throughput and satisfaction.
These operational requirements and the higher-acuity case mix we routinely manage are not represented in the payer’s QPA dataset and therefore warrant an upward deviation. The Qualified Payment Amount (QPA) is derived from median in-network rates that often blend hospital outpatient and urgent-care data—entities that do not share our 24/7 readiness, staffing ratios, or advanced clinical scope. Moreover, as a community-based emergency care provider, our case mix includes a broad range of high-acuity and after-hours presentations that cannot safely be managed by “ordinary providers” or urgent care facilities.
Accordingly, under 45 C.F.R. § 149.510(c)(4)(iii)(C), the certified IDR entity must give substantial weight to the provider’s scope of services, case mix, and clinical capabilities when determining the appropriate out-of-network rate. Our facility’s distinct operational and clinical profile materially differentiates us from the entities included in BCBS’s QPA calculation and establishes that the payer’s presumptive rate fails to capture the true cost and complexity of emergency care delivered in this setting.
This evidence affirms the critical role of **{billing_provider}** in providing high-acuity, hospital-level emergency services that uphold national quality standards and advance the protections intended under the No Surprises Act, ensuring that reimbursement determinations reflect both fairness and the indispensable public-health function of FSEDs

### Teaching Status, Case Mix, & Scope of Services
As a freestanding emergency center, **{billing_provider}** provides comprehensive emergency care across a broad and complex case mix. Our facility manages adult and pediatric emergencies, including critical conditions such as stroke and trauma stabilization, acute abdominal, cardiac complaints, and mental health evaluations. This diverse case mix reflects the high acuity and clinical complexity our physicians address daily.
Our scope of services is extensive and aligns with the full continuum of emergency care, including but not limited to:
•	Advanced diagnostic imaging (multi-slice CT scans, digital radiography)
•	Comprehensive laboratory diagnostics (CBC, BMP, LFTs, cardiac biomarkers)
•	Emergency medication administration, including critical care drugs
•	Intravenous fluid therapy and resuscitation
•	Short-term patient observation and stabilization
•	Coordinated emergency discharge planning and seamless transitions to higher levels of care or outpatient follow-up
Our board-certified emergency physicians and highly trained clinical staff maintain readiness to manage high-acuity patients around the clock, ensuring rapid response and quality outcomes. This case exemplifies the breadth and depth of our clinical capabilities and the critical nature of services provided by a Texas FSED.

### Market Share Considerations
In our region, we serve as a critical access point for emergency care, especially in after-hours and weekend settings where hospital ERs are overwhelmed. As a nonparticipating provider, our patient volume is modest compared to the regional reach of large insurers like BCBS. BCBS’s dominant market presence provides them disproportionate leverage in contract negotiations and QPA calculations—contributing to distorted reimbursement benchmarks that do not reflect actual cost or complexity.

### History of Network Negotiations and Good Faith Efforts
Over the past several years, we have engaged Blue Cross Blue Shield (BCBS) in multiple efforts to establish a fair and sustainable network agreement. Despite these initiatives, BCBS has either remained unresponsive or proposed contract terms that fall well below reasonable and sustainable reimbursement thresholds. Accepting such rates would directly undermine our ability to maintain 24/7 physician coverage by board-certified emergency medicine providers and to continue investing in the advanced diagnostic and clinical resources required for emergency care delivery.
As a freestanding emergency department, our statutory duty under the Prudent Layperson Standard is to provide emergency medical care to any patient presenting with symptoms of a potential emergency—without regard to insurance status or network participation. We are not permitted to defer or deny care based on contractual considerations. The lack of a network agreement, therefore, reflects not a refusal by the provider to participate, but rather BCBS’s use of market dominance to impose unsustainable contract terms. Under 45 C.F.R. § 149.510(c)(4)(iii)(B), credible evidence of such contracting history requires the certified IDR entity to consider whether the plan’s conduct has distorted the Qualified Payment Amount (QPA). Because BCBS’s self-reported QPA is derived from an in-network dataset shaped by its own suppressed contract rates, the presumption of QPA accuracy should not apply. Accordingly, this factor supports an upward adjustment to the QPA to reflect fair market value for the emergency services rendered.

### Other 
Under 45 C.F.R. § 149.510(c)(4)(iii), the plan’s Qualifying Payment Amount (QPA)—defined as the plan-calculated median in-network rate from 2019, adjusted for inflation—is presumed to represent the appropriate out-of-network (OON) rate. However, this presumption is fundamentally flawed when applied to freestanding emergency departments (FSEDs) due to key methodological and market limitations inherent in the payer’s QPA calculation.
Most FSEDs operate as non-contracted facilities, resulting in minimal or no in-network data from which a valid median contracted rate can be derived. Consequently, payers often substitute unrelated facility data—such as urgent care centers, physician offices, or hospital outpatient departments—to estimate QPAs. This approach violates 45 C.F.R. § 149.140(a)(8), which mandates that QPAs be based on “similar items and services” furnished by providers in the same or comparable specialty and facility type.
Payers frequently calculate and report QPAs based on self-reported, unaudited internal data, lacking external verification and methodological transparency, in breach of the disclosure requirements under 45 C.F.R. § 149.140(a)(12). This opacity obstructs providers’ ability to meaningfully evaluate the accuracy or fairness of proposed payments, thereby impairing the good-faith negotiation process envisioned by the No Surprises Act.
FSEDs incur substantial 24/7 operational and clinical readiness costs equivalent to those of hospital-based emergency departments—such as maintaining board-certified emergency physicians, on-site advanced imaging (CT, ultrasound, digital radiography), full laboratory services, intravenous medication administration, and critical-care stabilization capabilities. These significant fixed and standby expenses are inherently excluded from payer QPA algorithms, which primarily rely on contracted rates for lower-acuity or outpatient care settings.
Therefore, while payers assert their QPA reflects a “fair” OON rate, the underlying data are incomplete, non-comparable, and unverifiable, leading to a QPA that materially misrepresents the true cost and complexity of care delivered by FSEDs. Pursuant to § 149.510(c)(4)(iv), the provider submits credible, facility-specific evidence demonstrating that the QPA diverges materially from the appropriate market rate, including:
– Documentation of the training and board certification of attending emergency physicians;
– Evidence of the acuity and complexity inherent in emergency encounters routinely managed at the facility;
– Description of the scope and availability of emergency diagnostic and treatment resources;
– Verification of the facility’s continuous operational readiness and higher fixed costs; and
– Independent FAIR Health 80th-percentile benchmark data for ZIP code [insert ZIP], corroborating regional commercial rates considerably above the plan’s stated QPA.
None of this evidence relies on prohibited factors under § 149.510(c)(4)(v)—such as billed charges, usual and customary charges, or public payor rates. Instead, these data collectively illustrate that the payer’s QPA fails to capture the genuine cost and complexity of freestanding emergency care. Accordingly, the provider’s submitted rate represents the most accurate, market-reflective reimbursement aligned with the statutory intent and payment-determination framework of the No Surprises Act.

### Conclusion
In summary, the evidence clearly demonstrates that Blue Cross Blue Shield’s payment does not reflect the actual complexity or cost of the emergency services rendered, nor does it comply with the transparency and fairness standards established under the No Surprises Act and 45 C.F.R. §§149.140 and 149.510. By failing to disclose per-service QPAs and relying on internal, non-verifiable methodologies, BCBS has deprived the provider of the ability to assess payment accuracy or negotiate in good faith—contrary to both the spirit and letter of federal law.

**{billing_provider}** has consistently acted in good faith, delivering board-certified, 24/7 emergency care that meets nationally recognized clinical and quality standards. The services in this case were medically necessary, appropriately coded, and supported by full documentation. BCBS’s reliance on a “case-rate” framework misrepresents the nature of outpatient emergency billing and undermines equitable reimbursement practices.
We therefore respectfully request that the certified IDR entity issue a determination in favor of the provider’s offer. Such a decision would uphold the statutory payment-determination factors under 45 C.F.R. §149.510(c)(4)(iii), reinforce transparency in payer conduct, and preserve fair access to emergency medical care within the community.


""".strip()


# -----------------------------
# DOCX CREATION FUNCTIONS
# -----------------------------
def parse_bold_segments(paragraph, text):
    """Handle bold (**text**) and remove stray asterisks."""
    text = re.sub(r"\*\*([A-Za-z0-9\s:/()\-]+)\*+", r"**\1**", text)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2].strip())
            run.bold = True
        else:
            paragraph.add_run(part.replace("*", ""))


def add_formatted_paragraph(doc, text):
    """Handle headings, bullets, and normal paragraphs."""
    s = text.strip()
    if not s:
        doc.add_paragraph("")
        return

    # Headings
    if s.startswith("### ") or re.match(r"^[A-Z][A-Za-z\s&]+:$", s):
        heading = doc.add_paragraph()
        run = heading.add_run(s.replace("### ", "").strip(":").strip())
        run.bold = True
        run.underline = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(13)
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(6)
        return

    # Bulleted lists
    if s.startswith("•") or s.startswith("* "):
        clean_text = s.replace("•", "").replace("* ", "").strip()
        p = doc.add_paragraph(style="List Bullet")
        parse_bold_segments(p, clean_text)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        return

    # Regular text
    p = doc.add_paragraph()
    parse_bold_segments(p, s)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)


def create_docx_with_full_letter(full_letter):
    """Generate final DOCX with Aptos (Body) font."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Aptos (Body)"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos (Body)")
    style.font.size = Pt(12)

    for line in full_letter.split("\n"):
        add_formatted_paragraph(doc, line)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# -----------------------------
# MAIN WORKFLOW
# -----------------------------
if st.button("🚀 Run"):
    if not (eob_file and mrn_file and prompt_file):
        st.error("Please upload all required files.")
    else:
        with st.spinner("Processing... Please wait..."):
            try:
                eob_text = extract_text_from_pdf(eob_file)
                mrn_text = extract_text_from_pdf(mrn_file)
                prompt_text = prompt_file.read().decode("utf-8")

                date_of_service, hcpcs_codes, drg_code, billing_provider = extract_fields(eob_text)
                mrn_summary = generate_mrn_summary(prompt_text, mrn_text)
                full_letter = generate_bcbs_justification_letter(date_of_service, hcpcs_codes, drg_code, billing_provider, mrn_summary)

                output_doc = create_docx_with_full_letter(full_letter)

                st.success("✅ Automation complete!")
                st.subheader("📜 Generated BCBS Justification Letter")
                st.markdown(f"<div style='white-space: pre-wrap; font-size:16px;'>{full_letter}</div>", unsafe_allow_html=True)
                st.session_state["download_ready"] = True

                st.download_button(
                    label="📥 Download as Word Document (.docx)",
                    data=output_doc,
                    file_name="BCBS_Justification_for_IDR.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="download_docx_button",
                    help="Download the generated Word document (remains available after click).",
                )

            except Exception as e:
                st.error(f"❌ Error occurred: {str(e)}")